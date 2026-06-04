import os
import sys

def install_and_generate():
    print("正在为您检测并安装 Word 文档生成库 (python-docx)...")
    try:
        import docx
    except ImportError:
        import subprocess
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
        import docx

    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    print("正在生成高品质项目汇报 Word 文档...")
    doc = Document()

    # 设置页面页边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # 调色盘
    COLOR_PRIMARY = RGBColor(0, 102, 204)  # 苹果蓝
    COLOR_SECONDARY = RGBColor(28, 28, 30)  # 深黑
    COLOR_GRAY = RGBColor(120, 120, 128)  # 灰色

    # 1. 标题
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("《简阅》轻量级云原生小说阅读系统\n运维与故障自愈项目汇报书")
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = COLOR_PRIMARY
    title_run.font.name = "Microsoft YaHei"

    doc.add_paragraph().paragraph_format.space_after = Pt(24)

    # 2. 基本信息卡片
    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_format = info_p.paragraph_format
    p_format.line_spacing = 1.5
    
    runs_info = [
        ("项目名称：", True), ("《简阅（Yue）》轻量级小说平台\n", False),
        ("系统架构：", True), ("FastAPI + SQLAlchemy + SQLite + Docker Compose + Vue 3\n", False),
        ("部署环境：", True), ("openEuler 虚拟机 (IP: 10.42.78.219)\n", False),
        ("安全策略：", True), ("三级防御机制（脚本级自愈、热数据容灾、系统级快照）\n", False),
        ("生成时间：", True), ("2026年6月4日", False)
    ]
    for text, is_bold in runs_info:
        r = info_p.add_run(text)
        r.font.name = "Microsoft YaHei"
        r.font.size = Pt(11)
        r.font.bold = is_bold
        if is_bold:
            r.font.color.rgb = COLOR_SECONDARY

    doc.add_page_break()

    # 3. 目录与系统功能介绍
    h1 = doc.add_heading(level=1)
    h1_run = h1.add_run("一、 系统功能与架构亮点")
    h1_run.font.bold = True
    h1_run.font.size = Pt(16)
    h1_run.font.color.rgb = COLOR_PRIMARY

    features = [
        ("🌟 零配置自适应部署：", "前端网络请求地址（API_BASE_URL）完全自适应。代码通过浏览器窗口位置 `window.location` 动态匹配云服务器 IP，杜绝因服务重启、域名修改带来的“数据加载异常”，实现开箱即用。"),
        ("📊 实时控制台数据大屏：", "管理员控制面板全新集成“系统运行数据大屏”，服务端利用 SQLAlchemy 聚合函数（`func.sum` 及 `group_by`）毫秒级实时统计全站小说总数、总章节数、全站点击热度及分类题材占比。"),
        ("🔍 首页模糊检索与黄金高亮：", "引入响应式检索流。首页顶部配备苹果原生风格搜索框，可对书名、作者、品类进行模糊查询，并使用正则表达式安全高亮工具（`highlightText`）将匹配字眼在卡片及 3D 书脊上渲染为醒目的黄金圆角底色。"),
        ("🗑️ 原地级联安全删除：", "管理员点击“编辑信息”直接唤醒独立大弹窗。新增“彻底删除”底层端点，级联物理清除书籍，自动干净移除其下所属所有章节、读者点评及书架记录，保持 SQLite 轻量数据库零垃圾残留。")
    ]

    for title, desc in features:
        p = doc.add_paragraph(style='List Bullet')
        p_format = p.paragraph_format
        p_format.space_before = Pt(6)
        p_format.space_after = Pt(6)
        p_format.line_spacing = 1.25
        
        r_title = p.add_run(title)
        r_title.font.bold = True
        r_title.font.size = Pt(10.5)
        r_title.font.color.rgb = COLOR_SECONDARY
        
        r_desc = p.add_run(desc)
        r_desc.font.size = Pt(10.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # 4. 运维操作指导（核心部分）
    h2 = doc.add_heading(level=1)
    h2_run = h2.add_run("二、 云原生运维与高可用实操清单")
    h2_run.font.bold = True
    h2_run.font.size = Pt(16)
    h2_run.font.color.rgb = COLOR_PRIMARY

    steps = [
        ("1. 容器监控与自愈", "命令：`/root/Yue/ops/monitor_yue.sh`", 
         "说明：监控脚本通过 curl 主动访问 /health 判定状态。如果服务故障，脚本会触发【Auto-Heal 自动修复闭环】，自动重启容器并在3秒后复检，最后输出诊断报告。同时展现 `docker ps` 与磁盘开销。\n【在此处插入：执行自愈巡检脚本的终端日志输出截图】"),
        ("2. 实时日志审计", "命令：`docker logs -f yue-app` (实时查看) 或 `docker logs --tail 200 yue-app`", 
         "说明：实时展示 FastAPI 后端 API 调用流水及调试日志，证明服务运行通畅，可作为错误审计依据。\n【在此处插入：docker logs 的日志日志截图片段】"),
        ("3. 容灾热备份", "命令：`/root/Yue/ops/backup_yue.sh`", 
         "说明：对活动数据库 `yue.db` 和整个项目结构进行快照式压缩打包，归档在 `/root/backups/yue` 下，生成带有时戳的容灾包。\n【在此处插入：备份脚本运行输出与 ls -lh 生成备份文件的截图】"),
        ("4. 数据库一键灾后恢复", "命令：`/root/Yue/ops/restore_yue_db.sh /root/backups/yue/备份名称.db`", 
         "说明：模拟数据库物理损坏时的一键覆盖与容器自愈。恢复完成后自动调用健康检查确认 OK。\n【在此处插入：数据库一键恢复终端日志及 curl 200 的截图】"),
        ("5. 应用自动化部署", "前提：本地打包 Yue.zip 并上传到云服务器。\n命令：`/root/Yue/ops/deploy_yue.sh`", 
         "说明：脚本完成解压、权限设定、老旧容器强制拉下以及无缓存镜像强力重新编译拉起的全自动化流程。\n【在此处插入：自动化部署脚本运行成功后的终端截图】"),
        ("6. 访问与 API 吞吐验证", "验证方式：浏览器访问 http://10.42.78.219:8000 ；API 访问 http://10.42.78.219:8000/api/books", 
         "说明：验证前端功能是否全自适应显示，API JSON 数据是否通畅，证明部署成功。\n【在此处插入：浏览器主页面和 API 接口返回 JSON 数据的截图】")
    ]

    for title, cmd, desc in steps:
        h3 = doc.add_heading(level=2)
        h3_run = h3.add_run(title)
        h3_run.font.bold = True
        h3_run.font.size = Pt(12)
        h3_run.font.color.rgb = COLOR_SECONDARY
        
        p_cmd = doc.add_paragraph()
        p_cmd.paragraph_format.left_indent = Inches(0.2)
        r_cmd_label = p_cmd.add_run("🛠️ 操作指令：")
        r_cmd_label.font.bold = True
        r_cmd_val = p_cmd.add_run(cmd)
        r_cmd_val.font.bold = True
        r_cmd_val.font.color.rgb = COLOR_PRIMARY
        
        p_desc = doc.add_paragraph()
        p_desc.paragraph_format.left_indent = Inches(0.2)
        r_desc_label = p_desc.add_run("📝 说明及截图位置：\n")
        r_desc_label.font.bold = True
        r_desc_val = p_desc.add_run(desc)
        r_desc_val.font.italic = "【" in desc
        if "【" in desc:
            r_desc_val.font.color.rgb = RGBColor(220, 53, 69) # 红色警示
            
        doc.add_paragraph() # 空行隔开

    doc.add_page_break()

    # 5. 总结页
    h5 = doc.add_heading(level=1)
    h5_run = h5.add_run("三、 运维总结与项目收益说明")
    h5_run.font.bold = True
    h5_run.font.size = Pt(16)
    h5_run.font.color.rgb = COLOR_PRIMARY

    conclusion_p = doc.add_paragraph()
    conclusion_p.paragraph_format.line_spacing = 1.5
    conclusion_run = conclusion_p.add_run(
        "本项目成功在 openEuler 虚拟机上完成了轻量化小说阅读系统的生产级部署。在研发侧，项目攻克了“IP 频繁变动”这一分布式系统的经典痛点，设计了纯前端零配置自适应路由，极大降低了环境迁移难度；通过高维数据大屏、关键词毫秒级正则高亮以及级联删除，完美升级了管理端体验。\n\n"
        "在运维侧，引入了云原生理念。通过编写高自愈等级的巡检脚本 monitor_yue.sh，并配置定时任务 crontab，实现了服务宕机的“自动秒级发现与拉起自愈”；通过自动热备份及一键恢复脚本，搭建起完整的数据安全堡垒。整套方案实现了从“基础设施层、应用软件层、到持久化数据层”的多维度容灾监控，达到了真正的高可用状态。"
    )
    conclusion_run.font.size = Pt(10.5)

    output_path = os.path.join(os.path.dirname(__file__), "Yue_Deployment_Report.docx")
    doc.save(output_path)
    print(f"🎉 恭喜！高品质项目汇报 Word 文档已生成在：{output_path}")

if __name__ == '__main__':
    install_and_generate()
